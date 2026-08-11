"""DOCX document loader using python-docx. DOCX has no native page concept, so the whole
document is returned as a single logical page."""

import io

import docx

from app.ingestion.base import LoadedDocument, LoadedPage


class DocxLoader:
    """Loads text from Word (.docx) files."""

    def load(self, data: bytes, filename: str) -> LoadedDocument:
        """Extract paragraph text from the given DOCX bytes."""
        document = docx.Document(io.BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return LoadedDocument(
            source=filename,
            content_type=content_type,
            pages=[LoadedPage(text=text, page_number=None)],
        )
